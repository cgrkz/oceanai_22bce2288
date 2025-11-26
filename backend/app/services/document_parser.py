"""
Document Parser Service - Extracts text from various document formats.
Supports: MD, TXT, JSON, PDF, HTML, DOCX, DOC
"""

import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from backend.app.utils.logger import init_logger

# Initialize logger
logger = init_logger()


class DocumentChunk:
    """Represents a chunk of document text with metadata"""

    def __init__(
        self,
        text: str,
        metadata: Dict[str, Any],
        chunk_id: Optional[int] = None
    ):
        self.text = text
        self.metadata = metadata
        self.chunk_id = chunk_id

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "text": self.text,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id
        }


class DocumentParser:
    """Parse and extract text from various document formats"""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """
        Initialize document parser

        Args:
            chunk_size: Maximum size of text chunks
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        logger.info(
            "DocumentParser initialized",
            extra={
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap
            }
        )

    def parse_file(
        self,
        file_path: str,
        file_type: Optional[str] = None
    ) -> List[DocumentChunk]:
        """
        Parse a file and extract text chunks

        Args:
            file_path: Path to the file
            file_type: File extension (e.g., '.md', '.pdf')

        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Parsing file: {file_path}")

        try:
            path = Path(file_path)

            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            # Determine file type
            if file_type is None:
                file_type = path.suffix.lower()

            # Extract text based on file type
            if file_type in ['.md', '.txt']:
                text = self._parse_text_file(path)
            elif file_type == '.json':
                text = self._parse_json_file(path)
            elif file_type == '.pdf':
                text = self._parse_pdf_file(path)
            elif file_type == '.html':
                text = self._parse_html_file(path)
            elif file_type in ['.docx', '.doc']:
                text = self._parse_docx_file(path)
            else:
                logger.warning(f"Unsupported file type: {file_type}, treating as text")
                text = self._parse_text_file(path)

            # Create metadata
            metadata = {
                "source_document": path.name,
                "file_type": file_type,
                "file_path": str(path)
            }

            # Chunk the text
            chunks = self._create_chunks(text, metadata)

            logger.log_document_processing(
                filename=path.name,
                file_type=file_type,
                status="success",
                chunks=len(chunks)
            )

            return chunks

        except Exception as e:
            logger.log_document_processing(
                filename=Path(file_path).name if file_path else "unknown",
                file_type=file_type or "unknown",
                status="failed",
                error=str(e)
            )
            raise

    def _parse_text_file(self, path: Path) -> str:
        """Parse plain text or markdown file"""
        logger.debug(f"Parsing text file: {path.name}")

        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()

            logger.debug(f"Extracted {len(text)} characters from {path.name}")
            return text

        except UnicodeDecodeError:
            # Try with different encoding
            logger.warning(f"UTF-8 decoding failed for {path.name}, trying latin-1")
            with open(path, 'r', encoding='latin-1') as f:
                text = f.read()
            return text

    def _parse_json_file(self, path: Path) -> str:
        """Parse JSON file and convert to readable text"""
        logger.debug(f"Parsing JSON file: {path.name}")

        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert JSON to formatted string
        text = json.dumps(data, indent=2)

        logger.debug(f"Extracted {len(text)} characters from JSON: {path.name}")
        return text

    def _parse_pdf_file(self, path: Path) -> str:
        """Parse PDF file using PyMuPDF"""
        logger.debug(f"Parsing PDF file: {path.name}")

        try:
            doc = fitz.open(path)
            text = ""

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"

            doc.close()

            logger.debug(
                f"Extracted {len(text)} characters from {len(doc)} pages in {path.name}"
            )
            return text

        except Exception as e:
            logger.error(f"Error parsing PDF {path.name}: {str(e)}")
            raise

    def _parse_html_file(self, path: Path) -> str:
        """Parse HTML file"""
        logger.debug(f"Parsing HTML file: {path.name}")

        try:
            with open(path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Simple HTML tag removal (for basic extraction)
            # For production, consider using BeautifulSoup
            import re
            # Remove script and style elements
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)

            # Keep the HTML structure for context (important for testing)
            text = html_content

            logger.debug(f"Extracted {len(text)} characters from HTML: {path.name}")
            return text

        except Exception as e:
            logger.error(f"Error parsing HTML {path.name}: {str(e)}")
            raise

    def _parse_docx_file(self, path: Path) -> str:
        """Parse DOCX file"""
        logger.debug(f"Parsing DOCX file: {path.name}")

        try:
            doc = DocxDocument(path)
            text = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.debug(f"Extracted {len(text)} characters from DOCX: {path.name}")
            return text

        except Exception as e:
            logger.error(f"Error parsing DOCX {path.name}: {str(e)}")
            raise

    def _create_chunks(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks

        Args:
            text: Text to chunk
            metadata: Metadata for the document

        Returns:
            List of DocumentChunk objects
        """
        logger.debug(f"Creating chunks from text of length {len(text)}")

        if not text or len(text.strip()) == 0:
            logger.warning("Empty text provided for chunking")
            return []

        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size

            # If not the last chunk, try to break at sentence or word boundary
            if end < len(text):
                # Look for sentence boundary (., !, ?)
                sentence_end = max(
                    text.rfind('.', start, end),
                    text.rfind('!', start, end),
                    text.rfind('?', start, end)
                )

                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary (space)
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos

            # Extract chunk
            chunk_text = text[start:end].strip()

            if chunk_text:
                # Add chunk metadata
                chunk_metadata = metadata.copy()
                chunk_metadata['chunk_id'] = chunk_id
                chunk_metadata['chunk_size'] = len(chunk_text)

                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_id=chunk_id
                )

                chunks.append(chunk)
                chunk_id += 1

            # Move to next chunk with overlap
            start = end - self.chunk_overlap if end < len(text) else end

            # Prevent infinite loop
            if start >= len(text):
                break

        logger.debug(f"Created {len(chunks)} chunks")
        return chunks

    def parse_multiple_files(
        self,
        file_paths: List[str]
    ) -> List[DocumentChunk]:
        """
        Parse multiple files

        Args:
            file_paths: List of file paths

        Returns:
            Combined list of DocumentChunk objects
        """
        logger.info(f"Parsing {len(file_paths)} files")

        all_chunks = []

        for file_path in file_paths:
            try:
                chunks = self.parse_file(file_path)
                all_chunks.extend(chunks)
                logger.debug(f"Added {len(chunks)} chunks from {Path(file_path).name}")

            except Exception as e:
                logger.error(
                    f"Failed to parse {file_path}",
                    extra={"error": str(e)}
                )
                # Continue with other files

        logger.info(
            f"Parsed {len(file_paths)} files successfully",
            extra={"total_chunks": len(all_chunks)}
        )

        return all_chunks


# Create global instance
def get_document_parser() -> DocumentParser:
    """Get document parser instance"""
    from backend.config import settings

    return DocumentParser(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap
    )
